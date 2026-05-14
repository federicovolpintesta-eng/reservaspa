import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('Manual de Procedimientos: Recepción de Spa', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Los Pinos Resort & Spa', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('\n')

# 1. Acceso al Sistema
doc.add_heading('1. Acceso al Sistema (Panel de Administración)', level=1)
p = doc.add_paragraph()
p.add_run('URL de Acceso: ').bold = True
p.add_run('https://reservaspa.onrender.com/admin\n')
p.add_run('Contraseña: ').bold = True
p.add_run('spa2026\n\n')
doc.add_paragraph('Por motivos de seguridad, el acceso al administrador está oculto para los huéspedes. Siempre debe ingresar manualmente la palabra "/admin" al final de la dirección web. Al ingresar correctamente la contraseña, accederá al panel principal.')

# 2. Monitoreo de Reservas
doc.add_heading('2. Monitoreo Constante de Nuevas Reservas', level=1)
doc.add_paragraph('El sistema cuenta con un chequeo automático en segundo plano que busca nuevas reservas de huéspedes de forma constante.')
p = doc.add_paragraph()
p.add_run('Notificaciones del Navegador: ').bold = True
p.add_run('Es fundamental autorizar los "Permisos de Notificación" cuando el navegador de la recepción (Chrome, Safari, etc.) lo solicite. De esta forma, cada vez que un huésped finalice una reserva online, aparecerá una alerta visual y sonora en la esquina inferior derecha de la pantalla informando el nombre del huésped y el horario solicitado, sin necesidad de recargar la página.')

# 3. Grilla de Reservas (Kanban)
doc.add_heading('3. Gestión de la Grilla de Reservas (Kanban)', level=1)
doc.add_paragraph('El tablero Kanban permite visualizar los turnos de toda la semana (7 días) de forma panorámica. Cuenta con una barra de desplazamiento inferior para moverse por los diferentes días.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Arrastrar y Soltar (Drag & Drop): ').bold = True
p.add_run('Si un huésped solicita un cambio de horario o fecha, simplemente haga clic mantenido sobre la tarjeta de su reserva y arrástrela al nuevo día y bloque horario correspondiente. El sistema actualizará automáticamente la base de datos en la nube.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Eliminar Reservas: ').bold = True
p.add_run('Para cancelar un turno, haga clic en la "X" ubicada en la esquina superior derecha de la tarjeta de reserva. Se le pedirá una confirmación antes de eliminarla permanentemente.')

# 4. Bloqueo de Turnos
doc.add_heading('4. Bloqueo de Disponibilidad', level=1)
doc.add_paragraph('Si requiere bloquear un horario específico para mantenimiento, descanso del personal o una reserva externa especial:')
p = doc.add_paragraph(style='List Number')
p.add_run('Diríjase a la sección "Bloquear Turnos".')
p = doc.add_paragraph(style='List Number')
p.add_run('Seleccione la Fecha y el Horario exacto.')
p = doc.add_paragraph(style='List Number')
p.add_run('Haga clic en "Bloquear Turno".')
doc.add_paragraph('Los turnos bloqueados se mostrarán al huésped en la aplicación web bajo la etiqueta "Ocupado". Puede desbloquear el turno en cualquier momento desde la misma sección haciendo clic en la "X" roja junto al horario bloqueado.')

# 5. Gestión de Paquetes de Sesiones Múltiples
doc.add_heading('5. Gestión de Tratamientos Múltiples', level=1)
doc.add_paragraph('Cuando un huésped reserva online tratamientos complejos (ej. "Masaje Descontracturante — 5 Sesiones"), el sistema reacciona de la siguiente forma:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Agenda exclusivamente la primera sesión.\n')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Imprime una leyenda en el Comprobante (Voucher) del huésped indicando que el resto de las sesiones deben coordinarse en recepción.')
doc.add_paragraph('Procedimiento Operativo: Al recibir al huésped para su primera sesión, el recepcionista debe acordar de forma presencial los días y horarios de las sesiones restantes. Inmediatamente, se deben bloquear esos turnos manualmente en el sistema usando la herramienta "Bloquear Turnos" para garantizar el espacio y evitar la sobreventa online.')

# 6. Reportes y Métricas Financieras
doc.add_heading('6. Generación de Reportes y Métricas', level=1)
doc.add_paragraph('El sistema calcula automáticamente la ocupación, facturación total, ticket promedio, servicio más vendido y horario pico.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Filtro por Rango de Fechas: ').bold = True
p.add_run('Utilice los calendarios "Desde" y "Hasta" para limitar las métricas a un fin de semana específico, un mes o una temporada alta. Las estadísticas se recalcularán instantáneamente.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Impresión de Reportes: ').bold = True
p.add_run('Haciendo clic en el botón "🖨️ Imprimir Reporte", el sistema ocultará automáticamente el Kanban, los filtros y los botones innecesarios para generar un documento limpio (A4). Este reporte es ideal para controles de caja diarios o rendiciones gerenciales, y puede imprimirse en papel o guardarse como PDF.')

# 7. Gestión del Catálogo (Precios y Fotos)
doc.add_heading('7. Gestión del Catálogo (Precios y Fotos)', level=1)
doc.add_paragraph('Para mantener el menú de servicios actualizado sin necesidad de programadores:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Actualizar Precios: ').bold = True
p.add_run('En la sección "Gestión de Servicios", puede editar el precio de cualquier tratamiento. Los cambios se reflejarán inmediatamente para los huéspedes.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Subir Fotos: ').bold = True
p.add_run('Al crear o editar un servicio, puede seleccionar una foto desde su computadora. El sistema la optimizará automáticamente en la nube (ImgBB) y la mostrará en la galería de reservas del huésped.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Eliminar Servicios: ').bold = True
p.add_run('Use el icono del basurero (🗑️) para quitar tratamientos que ya no estén disponibles en la temporada.')

doc.save('/Users/federicovolpintesta/Desktop/Procedimiento_Recepcion_Spa.docx')
print("Documento Word generado exitosamente.")
